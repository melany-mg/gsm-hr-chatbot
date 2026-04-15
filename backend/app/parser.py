import pdfplumber
from docx import Document
from pathlib import Path
from typing import Generator


def _serialize_table(table: list[list]) -> str:
    """Serialize a table so each data row carries its column headers as context.

    The first row is treated as the header.  Every subsequent row becomes a
    comma-separated key:value sentence, making each row semantically
    self-contained for embedding.

    Cell content is flattened to a single line (internal newlines → space) so
    that multi-line PDF headers like "Years of employment\\nfrom date of hire"
    don't produce ambiguous line breaks that confuse the LLM.

    Example output for the PTO accrual table:
        Years of employment from date of hire | Equivalent days per year | ...
        Years of employment from date of hire: Less than 2, Equivalent days per year: 10 Days, ...
        Years of employment from date of hire: 2 through 4, Equivalent days per year: 15 Days, ...
    """
    if not table:
        return ""

    def clean(cell) -> str:
        if cell is None:
            return ""
        return " ".join(cell.split())  # collapse all whitespace including \n

    header = [clean(cell) for cell in table[0]]
    rows = [" | ".join(header)]  # keep plain header row for overall context
    for row in table[1:]:
        cells = [clean(cell) for cell in row]
        pairs = [
            f"{h}: {c}"
            for h, c in zip(header, cells)
            if h or c
        ]
        if pairs:
            rows.append(", ".join(pairs))
    return "\n".join(rows)


def extract_from_pdf(path: Path) -> Generator[dict, None, None]:
    """Yield one dict per non-empty page: {text, page, source}.

    Tables are extracted once via extract_tables() and serialized with full
    column-header context per row.  The prose text is extracted separately
    from non-table regions so table cell content is never duplicated.
    """
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            parts = []

            # Identify table bounding boxes so we can exclude them from the
            # plain-text extraction (pdfplumber's extract_text() includes table
            # cell text, which would otherwise duplicate the structured data).
            tables_on_page = page.find_tables()
            table_bboxes = [t.bbox for t in tables_on_page]

            if table_bboxes:
                def outside_tables(obj, _bboxes=table_bboxes):
                    for bbox in _bboxes:
                        if (
                            obj.get("x0", 0) >= bbox[0] - 1
                            and obj.get("x1", 0) <= bbox[2] + 1
                            and obj.get("top", 0) >= bbox[1] - 1
                            and obj.get("bottom", 0) <= bbox[3] + 1
                        ):
                            return False
                    return True

                text = page.filter(outside_tables).extract_text() or ""
            else:
                text = page.extract_text() or ""

            if text.strip():
                parts.append(text.strip())

            for table in page.extract_tables():
                serialized = _serialize_table(table)
                if serialized:
                    parts.append(serialized)

            combined = "\n".join(parts).strip()
            if combined:
                yield {"text": combined, "page": i, "source": path.name}


def extract_from_docx(path: Path) -> Generator[dict, None, None]:
    """Yield a single dict for the whole document: {text, page, source}.
    DOCX files have no native page numbers so page is always 1."""
    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        serialized = _serialize_table(
            [[cell.text.strip() for cell in row.cells] for row in table.rows]
        )
        if serialized:
            parts.append(serialized)
    combined = "\n".join(parts).strip()
    if combined:
        yield {"text": combined, "page": 1, "source": path.name}
